from datetime import datetime
from docx import Document


def add_kv_table(doc: Document, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Topic"
    hdr[1].text = "Current Position"
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = k
        cells[1].text = v
    return table


def main():
    doc = Document()
    today = datetime.now().strftime("%Y-%m-%d")

    doc.add_heading("ROAM V3 Founder Briefing", level=0)
    doc.add_paragraph(
        f"Prepared: {today}\n"
        "Purpose: give founders a clear picture of product vision, current build status, "
        "design/PRD alignment, and execution path."
    )

    doc.add_heading("1) Executive Summary", level=1)
    doc.add_paragraph(
        "ROAM is being built as a choreography operating workspace, not a generic media app. "
        "The core promise is to reduce in-studio friction (looping, reference recall, section organization) "
        "and preserve creative memory across sessions."
    )
    doc.add_paragraph(
        "Current state: strong momentum and visible progress. Route-level coverage for key V3 pages exists, "
        "with ongoing parity work required to fully match Figma + PRD behavior standards."
    )

    doc.add_heading("2) Product Thesis (Founder Narrative)", level=1)
    doc.add_paragraph(
        "The problem: choreographers lose ideas because current tools are playback-first and memory-fragmented."
    )
    doc.add_paragraph(
        "The thesis: a session-native workspace (workbench, song-map, capture, reference, group) compounds value "
        "because it stores and structures creative decisions over time."
    )
    doc.add_paragraph(
        "Design principle: 'mirror, not muse' — ROAM shows the work clearly and avoids prescriptive/evaluative behavior."
    )

    doc.add_heading("3) Scope and Source of Truth", level=1)
    add_kv_table(
        doc,
        [
            ("UI/Interaction source", "Figma ROAMV3, canonical v3-locked set"),
            ("Behavior/Priority source", "ROAM_PRD_FINAL.md"),
            ("Implementation tracker", "V3_COVERAGE_AUDIT.md + V3_IMPLEMENTATION_BACKLOG.md + V3_PROGRESS_LOG.md"),
            ("Target devices", "Tablet landscape (iPad/Android tablet) + Android phone portrait"),
        ],
    )

    doc.add_heading("4) Current Delivery Status", level=1)
    doc.add_paragraph("Key implementation progress already completed:")
    for item in [
        "Song-map route baseline added",
        "Spatial route baseline added",
        "Group route baseline added (choreographer/dancer scaffolding)",
        "Workbench parity passes started (section swipe, transport presets, clip style polish)",
        "Figma v3-locked cleaned to a canonical frame set",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("Main gap: many pages are still Partial, not Implemented at strict acceptance.")

    doc.add_heading("5) Feature Decisions Locked (Important for Founders)", level=1)
    doc.add_paragraph("Music Partition:")
    for item in [
        "Must ship in V3 as an optional mode (not default).",
        "Opened via visible toggle (Counts | Partition).",
        "Read-only in V3; notation editing is out of scope.",
    ]:
        doc.add_paragraph(item, style="List Bullet 2")

    doc.add_paragraph("Dual-screen capture:")
    for item in [
        "Two-device workflow remains primary in V3.",
        "Single-device dual-screen is scoped optional beta behind explicit toggle.",
        "Never default-on; auto-fallback to standard capture if performance degrades.",
    ]:
        doc.add_paragraph(item, style="List Bullet 2")

    doc.add_paragraph("AI scope:")
    for item in [
        "No generative choreography in current V3 core.",
        "No evaluative movement scoring in V3 core.",
        "AI remains a later assistive extension if retention validates the need.",
    ]:
        doc.add_paragraph(item, style="List Bullet 2")

    doc.add_heading("6) What Is Needed to Reach 'Validation-Ready'", level=1)
    for item in [
        "Complete P0 parity for workbench/song-map/spatial/group interactions.",
        "Ref-viewer architecture parity (sheet behavior + loop/save flow).",
        "Phone UX hardening beyond capture/ref-viewer/group dancer paths.",
        "Close visual-system mismatches across home/auth/library/settings.",
        "End-to-end confidence checks on tablet + Android phone flows.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7) 30-60-90 Execution Plan", level=1)
    doc.add_paragraph("Next 30 days (stabilize core):", style="Intense Quote")
    for item in [
        "Finish high-impact workbench parity and interaction fit.",
        "Raise song-map/spatial/group from baseline to reliable core workflows.",
        "Finalize Figma-to-code parity matrix updates after each implementation batch.",
    ]:
        doc.add_paragraph(item, style="List Bullet 2")

    doc.add_paragraph("60 days (product integrity):", style="Intense Quote")
    for item in [
        "Ref-viewer, capture, library, auth, home, settings parity closure.",
        "Device-specific UX QA across tablet landscape and Android phone portrait.",
        "Operational hardening for session continuity and predictable behavior.",
    ]:
        doc.add_paragraph(item, style="List Bullet 2")

    doc.add_paragraph("90 days (founder-ready validation package):", style="Intense Quote")
    for item in [
        "PRD-to-implementation traceability completed.",
        "Pilot-ready build + usability outcomes from real users.",
        "Clear metrics narrative: retention assumptions vs observed behavior.",
    ]:
        doc.add_paragraph(item, style="List Bullet 2")

    doc.add_heading("8) Risks and Mitigations", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    h = table.rows[0].cells
    h[0].text = "Risk"
    h[1].text = "Impact"
    h[2].text = "Mitigation"
    risk_rows = [
        (
            "Scope creep from optional advanced modes",
            "Delays core parity and launch confidence",
            "Ship optional modes behind explicit toggles with strict acceptance gates",
        ),
        (
            "Device performance variance (Android phone/tablet)",
            "Inconsistent UX in studio context",
            "Performance-gated fallback behavior + dedicated device QA cycles",
        ),
        (
            "Visual polish without behavior closure",
            "Looks complete but fails in real choreography workflows",
            "Track PRD behavior completion separately from visual parity",
        ),
    ]
    for r in risk_rows:
        c = table.add_row().cells
        c[0].text, c[1].text, c[2].text = r

    doc.add_heading("9) Suggested Founder Pitch Track (3 minutes)", level=1)
    doc.add_paragraph("Use this sequence:", style="Intense Quote")
    for item in [
        "Start with the pain: 20-30 minutes lost per rehearsal to loop friction and context switching.",
        "State ROAM's role: session memory + structure + capture in one choreography workspace.",
        "Show what's already built: canonical design system, route coverage, and active parity execution.",
        "Clarify strategic choices: optional partition mode, scoped dual-screen beta, non-prescriptive AI stance.",
        "Close with roadmap confidence: clear P0/P1 plan to validation-ready state.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("10) Current Ask / Decision for Founders", level=1)
    doc.add_paragraph(
        "Approve the execution focus: finish P0 behavior + parity before expansion features, "
        "while keeping optional partition and dual-screen in scoped mode as defined."
    )

    output_path = r"c:\Users\Georges\Documents\Cursor 2 V3\docs\ROAM_Founder_Briefing_Report.docx"
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
